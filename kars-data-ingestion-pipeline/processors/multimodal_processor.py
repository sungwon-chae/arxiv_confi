"""
멀티모달 문서 프로세서
MinerU/magic-pdf 기반 고품질 문서 구조 추출 및 멀티모달 처리
"""

import os
import logging
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
from datetime import datetime
import tempfile
import base64
import json
from dataclasses import dataclass
import hashlib

# MinerU/magic-pdf 기반 문서 파싱
try:
    from magic_pdf.pdf_parse_by_txt import parse_pdf
    from magic_pdf.libs.MakeContentConfig import MakeMode
    MAGIC_PDF_AVAILABLE = True
except ImportError:
    MAGIC_PDF_AVAILABLE = False
    logging.warning("magic-pdf not available. Install with: pip install magic-pdf[full]")

# Marker PDF 대안
try:
    from marker.convert import convert_single_pdf
    from marker.models import load_all_models
    MARKER_AVAILABLE = True
except ImportError:
    MARKER_AVAILABLE = False
    logging.warning("marker-pdf not available. Install with: pip install marker-pdf")

# Fallback processors
import pypdf
import docx
from PIL import Image
import pandas as pd
import numpy as np

# OCR support
try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("pytesseract not available for OCR")

logger = logging.getLogger(__name__)

@dataclass
class ExtractedElement:
    """추출된 문서 요소"""
    element_type: str  # 'text', 'image', 'table', 'heading', 'equation', 'list'
    content: Union[str, bytes, Dict]
    metadata: Dict[str, Any]
    page_number: Optional[int] = None
    position: Optional[Dict[str, float]] = None  # x, y, width, height
    confidence: float = 1.0
    relationships: List[str] = None  # 다른 요소와의 관계

class MultimodalDocumentProcessor:
    """멀티모달 문서 처리기"""
    
    def __init__(self, 
                 use_gpu: bool = False,
                 ocr_enabled: bool = True,
                 preserve_layout: bool = True,
                 extract_equations: bool = True):
        self.use_gpu = use_gpu
        self.ocr_enabled = ocr_enabled and OCR_AVAILABLE
        self.preserve_layout = preserve_layout
        self.extract_equations = extract_equations
        
        # 처리 엔진 우선순위
        self.pdf_engine = self._detect_pdf_engine()
        
        # Marker 모델 초기화 (필요시)
        self.marker_models = None
        if self.pdf_engine == "marker" and MARKER_AVAILABLE:
            try:
                self.marker_models = load_all_models()
                logger.info("Marker PDF 모델 로드 완료")
            except Exception as e:
                logger.error(f"Marker 모델 로드 실패: {e}")
                self.pdf_engine = "fallback"
    
    def _detect_pdf_engine(self) -> str:
        """사용 가능한 PDF 처리 엔진 감지"""
        if MAGIC_PDF_AVAILABLE:
            return "magic_pdf"
        elif MARKER_AVAILABLE:
            return "marker"
        else:
            return "fallback"
    
    async def process_document(self, 
                             file_path: Union[str, Path],
                             extract_options: Optional[Dict[str, bool]] = None,
                             output_format: str = "structured_json") -> Dict[str, Any]:
        """
        문서 처리 메인 함수
        
        Args:
            file_path: 처리할 문서 파일 경로
            extract_options: 추출 옵션 {extract_text, extract_images, extract_tables}
            output_format: 출력 형식 (structured_json, markdown, elements)
            
        Returns:
            처리된 문서 데이터
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        # 기본 추출 옵션 설정
        if extract_options is None:
            extract_options = {
                "extract_text": True,
                "extract_images": True,
                "extract_tables": True,
                "preserve_layout": True
            }
        
        logger.info(f"문서 처리 시작: {file_path}")
        
        try:
            if file_path.suffix.lower() == '.pdf':
                if self.pdf_engine == "magic_pdf":
                    return await self._process_with_magic_pdf(file_path, extract_options, output_format)
                elif self.pdf_engine == "marker":
                    return await self._process_with_marker(file_path, extract_options, output_format)
                else:
                    return await self._process_with_fallback(file_path, extract_options, output_format)
            else:
                # Non-PDF documents
                return await self._process_with_fallback(file_path, extract_options, output_format)
        except Exception as e:
            logger.error(f"문서 처리 실패: {e}")
            return {
                "success": False,
                "error": str(e),
                "file_path": str(file_path),
                "timestamp": datetime.now().isoformat()
            }
    
    async def _process_with_magic_pdf(self, 
                                    file_path: Path, 
                                    extract_options: Dict[str, bool], 
                                    output_format: str) -> Dict[str, Any]:
        """MinerU/magic-pdf를 사용한 고품질 문서 처리"""
        
        try:
            # magic-pdf로 문서 파싱
            with open(file_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
            
            # 파싱 옵션 설정
            parse_method = MakeMode.LAYOUT_MODE if self.preserve_layout else MakeMode.TEXT_MODE
            
            # PDF 파싱 실행
            parsed_result = parse_pdf(
                pdf_bytes,
                parse_method=parse_method,
                model_list=['layout', 'formula', 'table'] if self.extract_equations else ['layout', 'table']
            )
            
            extracted_elements = []
            
            # 파싱된 결과에서 요소 추출
            for page_idx, page_data in enumerate(parsed_result.get('pages', [])):
                page_number = page_idx + 1
                
                # 텍스트 블록 처리
                if extract_options.get("extract_text", True):
                    for block in page_data.get('blocks', []):
                        if block['type'] == 'text':
                            element = ExtractedElement(
                                element_type="text",
                                content=block['text'],
                                metadata={
                                    "text_type": self._detect_text_type(block),
                                    "font_size": block.get('font_size'),
                                    "font_name": block.get('font_name'),
                                    "block_id": block.get('block_id')
                                },
                                page_number=page_number,
                                position={
                                    "x": block['bbox'][0],
                                    "y": block['bbox'][1],
                                    "width": block['bbox'][2] - block['bbox'][0],
                                    "height": block['bbox'][3] - block['bbox'][1]
                                },
                                confidence=block.get('confidence', 1.0)
                            )
                            extracted_elements.append(element)
                
                # 테이블 처리
                if extract_options.get("extract_tables", True):
                    for table in page_data.get('tables', []):
                        table_data = self._process_table_structure(table)
                        element = ExtractedElement(
                            element_type="table",
                            content=table_data,
                            metadata={
                                "rows": len(table_data['rows']),
                                "columns": len(table_data['rows'][0]) if table_data['rows'] else 0,
                                "has_header": table.get('has_header', False)
                            },
                            page_number=page_number,
                            position={
                                "x": table['bbox'][0],
                                "y": table['bbox'][1],
                                "width": table['bbox'][2] - table['bbox'][0],
                                "height": table['bbox'][3] - table['bbox'][1]
                            }
                        )
                        extracted_elements.append(element)
                
                # 이미지 처리
                if extract_options.get("extract_images", True):
                    for image in page_data.get('images', []):
                        img_element = await self._process_image_with_ocr(image, page_number)
                        if img_element:
                            extracted_elements.append(img_element)
                
                # 수식 처리
                if self.extract_equations and extract_options.get("extract_equations", True):
                    for equation in page_data.get('equations', []):
                        element = ExtractedElement(
                            element_type="equation",
                            content={
                                "latex": equation.get('latex', ''),
                                "text": equation.get('text', '')
                            },
                            metadata={
                                "inline": equation.get('inline', False),
                                "equation_id": equation.get('id')
                            },
                            page_number=page_number,
                            position=equation.get('bbox')
                        )
                        extracted_elements.append(element)
            
            # 요소 간 관계 분석
            extracted_elements = self._analyze_element_relationships(extracted_elements)
            
            return self._format_output(file_path, extracted_elements, output_format)
            
        except Exception as e:
            logger.error(f"magic-pdf 처리 실패: {e}")
            # Fallback to marker or basic processing
            if self.pdf_engine == "marker":
                return await self._process_with_marker(file_path, extract_options, output_format)
            else:
                return await self._process_with_fallback(file_path, extract_options, output_format)
    
    def _detect_text_type(self, block: Dict) -> str:
        """텍스트 블록의 타입 감지"""
        # 폰트 크기와 스타일 기반 감지
        font_size = block.get('font_size', 12)
        font_name = block.get('font_name', '').lower()
        text = block.get('text', '')
        
        if font_size > 16 or 'bold' in font_name:
            if len(text) < 100:
                return 'heading'
        
        if block.get('is_list_item'):
            return 'list_item'
        
        return 'paragraph'
    
    def _process_table_structure(self, table: Dict) -> Dict:
        """테이블 구조 처리 및 정리"""
        rows = []
        for row in table.get('rows', []):
            row_data = []
            for cell in row.get('cells', []):
                row_data.append(cell.get('text', ''))
            rows.append(row_data)
        
        return {
            "rows": rows,
            "headers": table.get('headers', []),
            "structure": table.get('structure', {})
        }
    
    async def _process_image_with_ocr(self, image_data: Dict, page_number: int) -> Optional[ExtractedElement]:
        """이미지 처리 및 OCR 수행"""
        try:
            img_bytes = image_data.get('data')
            if not img_bytes:
                return None
            
            # PIL Image로 변환
            img = Image.open(base64.b64decode(img_bytes))
            
            ocr_text = ""
            if self.ocr_enabled:
                try:
                    # OCR 수행
                    ocr_text = pytesseract.image_to_string(img)
                except Exception as e:
                    logger.warning(f"OCR 실패: {e}")
            
            # 이미지 메타데이터 추출
            metadata = {
                "format": img.format,
                "mode": img.mode,
                "width": img.width,
                "height": img.height,
                "has_text": bool(ocr_text.strip()),
                "ocr_text": ocr_text if ocr_text.strip() else None
            }
            
            return ExtractedElement(
                element_type="image",
                content=base64.b64encode(img_bytes).decode('utf-8'),
                metadata=metadata,
                page_number=page_number,
                position=image_data.get('bbox'),
                confidence=0.9 if ocr_text else 1.0
            )
            
        except Exception as e:
            logger.error(f"이미지 처리 실패: {e}")
            return None
    
    def _analyze_element_relationships(self, elements: List[ExtractedElement]) -> List[ExtractedElement]:
        """요소 간 공간적/의미적 관계 분석"""
        # 페이지별로 요소 그룹화
        page_elements = {}
        for elem in elements:
            page = elem.page_number or 0
            if page not in page_elements:
                page_elements[page] = []
            page_elements[page].append(elem)
        
        # 각 페이지에서 관계 분석
        for page, page_elems in page_elements.items():
            for i, elem1 in enumerate(page_elems):
                if not elem1.relationships:
                    elem1.relationships = []
                
                for j, elem2 in enumerate(page_elems):
                    if i != j and elem1.position and elem2.position:
                        # 공간적 근접성 확인
                        if self._are_elements_adjacent(elem1.position, elem2.position):
                            relationship = f"adjacent_to_{elem2.element_type}_{j}"
                            elem1.relationships.append(relationship)
                        
                        # 테이블과 캡션 관계
                        if elem1.element_type == "table" and elem2.element_type == "text":
                            if self._is_caption(elem2, elem1):
                                elem1.relationships.append(f"caption_{j}")
                                elem2.relationships = elem2.relationships or []
                                elem2.relationships.append(f"caption_of_table_{i}")
        
        return elements
    
    def _are_elements_adjacent(self, pos1: Dict, pos2: Dict) -> bool:
        """두 요소가 인접한지 확인"""
        # 간단한 거리 기반 확인
        threshold = 50  # 픽셀
        
        # 중심점 계산
        center1_x = pos1['x'] + pos1['width'] / 2
        center1_y = pos1['y'] + pos1['height'] / 2
        center2_x = pos2['x'] + pos2['width'] / 2
        center2_y = pos2['y'] + pos2['height'] / 2
        
        distance = ((center2_x - center1_x) ** 2 + (center2_y - center1_y) ** 2) ** 0.5
        return distance < threshold
    
    def _is_caption(self, text_elem: ExtractedElement, table_elem: ExtractedElement) -> bool:
        """텍스트가 테이블의 캡션인지 확인"""
        if not text_elem.content or not isinstance(text_elem.content, str):
            return False
        
        # 캡션 패턴 확인
        caption_patterns = ['table', '표', 'figure', '그림']
        text_lower = text_elem.content.lower()
        
        has_pattern = any(pattern in text_lower for pattern in caption_patterns)
        
        # 위치 확인 (테이블 위 또는 아래)
        if text_elem.position and table_elem.position:
            text_y = text_elem.position['y']
            table_y = table_elem.position['y']
            table_height = table_elem.position['height']
            
            is_above = text_y < table_y and (table_y - text_y) < 100
            is_below = text_y > (table_y + table_height) and (text_y - (table_y + table_height)) < 100
            
            return has_pattern and (is_above or is_below)
        
        return False
    
    async def _process_with_marker(self, 
                                 file_path: Path, 
                                 extract_options: Dict[str, bool], 
                                 output_format: str) -> Dict[str, Any]:
        """Marker PDF를 사용한 문서 처리"""
        try:
            # Marker로 PDF 변환
            full_text, images, metadata = convert_single_pdf(
                str(file_path),
                self.marker_models,
                batch_multiplier=2 if self.use_gpu else 1
            )
            
            extracted_elements = []
            
            # 마크다운 텍스트를 구조화된 요소로 변환
            if extract_options.get("extract_text", True):
                elements_from_markdown = self._parse_markdown_structure(full_text)
                extracted_elements.extend(elements_from_markdown)
            
            # 이미지 처리
            if extract_options.get("extract_images", True) and images:
                for idx, img in enumerate(images):
                    element = ExtractedElement(
                        element_type="image",
                        content=base64.b64encode(img).decode('utf-8'),
                        metadata={"image_index": idx},
                        confidence=0.9
                    )
                    extracted_elements.append(element)
            
            return self._format_output(file_path, extracted_elements, output_format)
            
        except Exception as e:
            logger.error(f"Marker 처리 실패: {e}")
            return await self._process_with_fallback(file_path, extract_options, output_format)
    
    def _parse_markdown_structure(self, markdown_text: str) -> List[ExtractedElement]:
        """마크다운 텍스트를 구조화된 요소로 파싱"""
        elements = []
        lines = markdown_text.split('\n')
        
        current_section = []
        for line in lines:
            if line.startswith('#'):
                # 헤딩 처리
                if current_section:
                    # 이전 섹션 저장
                    elements.append(ExtractedElement(
                        element_type="text",
                        content='\n'.join(current_section),
                        metadata={"text_type": "paragraph"}
                    ))
                    current_section = []
                
                # 헤딩 레벨 추출
                heading_level = len(line.split(' ')[0])
                heading_text = line.lstrip('#').strip()
                
                elements.append(ExtractedElement(
                    element_type="text",
                    content=heading_text,
                    metadata={
                        "text_type": "heading",
                        "heading_level": heading_level
                    }
                ))
            elif line.startswith('|'):
                # 테이블 처리
                table_lines = [line]
                # 테이블 끝까지 수집
                # ... (테이블 파싱 로직)
            else:
                current_section.append(line)
        
        # 마지막 섹션 저장
        if current_section:
            elements.append(ExtractedElement(
                element_type="text",
                content='\n'.join(current_section),
                metadata={"text_type": "paragraph"}
            ))
        
        return elements
    
    async def _process_with_fallback(self, 
                                   file_path: Path, 
                                   extract_options: Dict[str, bool], 
                                   output_format: str) -> Dict[str, Any]:
        """Fallback 프로세서를 사용한 문서 처리"""
        
        extracted_elements = []
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                elements = await self._process_pdf_fallback(file_path, extract_options)
            elif file_ext in ['.docx', '.doc']:
                elements = await self._process_docx_fallback(file_path, extract_options)
            elif file_ext in ['.txt', '.md']:
                elements = await self._process_text_fallback(file_path, extract_options)
            else:
                raise ValueError(f"지원하지 않는 파일 형식: {file_ext}")
            
            extracted_elements.extend(elements)
            
        except Exception as e:
            logger.error(f"Fallback 처리 실패: {e}")
            raise
        
        return self._format_output(file_path, extracted_elements, output_format)
    
    async def _process_pdf_fallback(self, file_path: Path, extract_options: Dict[str, bool]) -> List[ExtractedElement]:
        """PDF 파일 fallback 처리"""
        elements = []
        
        if extract_options.get("extract_text", True):
            try:
                with open(file_path, 'rb') as file:
                    reader = pypdf.PdfReader(file)
                    
                    for page_num, page in enumerate(reader.pages, 1):
                        text = page.extract_text()
                        if text.strip():
                            elements.append(ExtractedElement(
                                element_type="text",
                                content=text,
                                metadata={
                                    "text_type": "page_content",
                                    "confidence": 0.8,
                                    "length": len(text)
                                },
                                page_number=page_num
                            ))
            except Exception as e:
                logger.error(f"PDF 텍스트 추출 실패: {e}")
        
        # PDF 이미지/테이블 추출은 복잡하므로 RAG-Anything 권장 메시지
        if extract_options.get("extract_images", True) or extract_options.get("extract_tables", True):
            elements.append(ExtractedElement(
                element_type="text",
                content="⚠️ PDF 이미지/테이블 추출을 위해서는 RAG-Anything 라이브러리가 필요합니다.",
                metadata={
                    "text_type": "system_message",
                    "confidence": 1.0,
                    "length": 50
                }
            ))
        
        return elements
    
    async def _process_docx_fallback(self, file_path: Path, extract_options: Dict[str, bool]) -> List[ExtractedElement]:
        """DOCX 파일 fallback 처리"""
        elements = []
        
        if extract_options.get("extract_text", True):
            try:
                doc = docx.Document(file_path)
                
                for i, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        # 스타일 기반으로 헤딩 감지
                        text_type = "heading" if paragraph.style.name.startswith('Heading') else "paragraph"
                        
                        elements.append(ExtractedElement(
                            element_type="text",
                            content=paragraph.text,
                            metadata={
                                "text_type": text_type,
                                "style": paragraph.style.name,
                                "confidence": 0.9,
                                "length": len(paragraph.text)
                            }
                        ))
                
                # 테이블 추출
                if extract_options.get("extract_tables", True):
                    for i, table in enumerate(doc.tables):
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text for cell in row.cells]
                            table_data.append(row_data)
                        
                        elements.append(ExtractedElement(
                            element_type="table",
                            content={"rows": table_data},
                            metadata={
                                "rows": len(table.rows),
                                "columns": len(table.columns) if table.rows else 0,
                                "table_type": "docx_table"
                            }
                        ))
                        
            except Exception as e:
                logger.error(f"DOCX 처리 실패: {e}")
        
        return elements
    
    async def _process_text_fallback(self, file_path: Path, extract_options: Dict[str, bool]) -> List[ExtractedElement]:
        """텍스트 파일 fallback 처리"""
        elements = []
        
        if extract_options.get("extract_text", True):
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    
                    elements.append(ExtractedElement(
                        element_type="text",
                        content=content,
                        metadata={
                            "text_type": "plain_text",
                            "confidence": 1.0,
                            "length": len(content),
                            "encoding": "utf-8"
                        }
                    ))
                    
            except Exception as e:
                logger.error(f"텍스트 파일 처리 실패: {e}")
        
        return elements
    
    def _format_output(self, 
                      file_path: Path, 
                      elements: List[ExtractedElement], 
                      output_format: str) -> Dict[str, Any]:
        """출력 형식에 따른 결과 포맷팅"""
        
        base_result = {
            "success": True,
            "file_path": str(file_path),
            "timestamp": datetime.now().isoformat(),
            "total_elements": len(elements),
            "processing_method": self.pdf_engine if hasattr(self, 'pdf_engine') else "fallback"
        }
        
        if output_format == "structured_json":
            base_result["elements"] = []
            for elem in elements:
                base_result["elements"].append({
                    "type": elem.element_type,
                    "content": elem.content,
                    "metadata": elem.metadata,
                    "page_number": elem.page_number,
                    "position": elem.position
                })
                
        elif output_format == "markdown":
            markdown_content = f"# {file_path.name}\n\n"
            
            for elem in elements:
                if elem.element_type == "text":
                    if elem.metadata.get("text_type") == "heading":
                        markdown_content += f"## {elem.content}\n\n"
                    else:
                        markdown_content += f"{elem.content}\n\n"
                elif elem.element_type == "table":
                    markdown_content += "### 표\n\n"
                    if isinstance(elem.content, dict) and "rows" in elem.content:
                        for row in elem.content["rows"]:
                            markdown_content += "| " + " | ".join(str(cell) for cell in row) + " |\n"
                    markdown_content += "\n"
                elif elem.element_type == "image":
                    markdown_content += f"![이미지](data:image/png;base64,{elem.content[:50]}...)\n\n"
            
            base_result["markdown_content"] = markdown_content
            
        elif output_format == "elements":
            base_result["text_elements"] = [e for e in elements if e.element_type == "text"]
            base_result["image_elements"] = [e for e in elements if e.element_type == "image"]
            base_result["table_elements"] = [e for e in elements if e.element_type == "table"]
        
        # 통계 정보 추가
        element_stats = {}
        for elem in elements:
            elem_type = elem.element_type
            element_stats[elem_type] = element_stats.get(elem_type, 0) + 1
        
        base_result["element_statistics"] = element_stats
        
        return base_result
    
    def extract_for_vectorization(self, processed_doc: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        벡터화를 위한 청크 추출 (멀티모달 지원)
        
        Args:
            processed_doc: process_document()의 결과
            
        Returns:
            벡터화 가능한 청크들의 리스트
        """
        chunks = []
        
        if not processed_doc.get("success", False):
            return chunks
        
        elements = processed_doc.get("elements", [])
        
        for i, elem in enumerate(elements):
            chunk = {
                "id": f"{Path(processed_doc['file_path']).stem}_{elem['type']}_{i}",
                "content": "",
                "metadata": {
                    "source_file": processed_doc["file_path"],
                    "element_type": elem["type"],
                    "page_number": elem.get("page_number"),
                    "processing_timestamp": processed_doc["timestamp"],
                    "confidence": elem.get("confidence", 1.0),
                    "relationships": elem.get("relationships", []),
                    **elem.get("metadata", {})
                }
            }
            
            if elem["type"] == "text":
                chunk["content"] = elem["content"]
                # 헤딩인 경우 특별 표시
                if elem.get("metadata", {}).get("text_type") == "heading":
                    chunk["content"] = f"[제목] {chunk['content']}"
                    
            elif elem["type"] == "table":
                # 테이블을 구조화된 텍스트로 변환
                if isinstance(elem["content"], dict) and "rows" in elem["content"]:
                    headers = elem["content"].get("headers", [])
                    rows = elem["content"]["rows"]
                    
                    # 헤더가 있으면 포함
                    if headers:
                        table_text = "표 헤더: " + " | ".join(headers) + "\n"
                        table_text += "\n".join([" | ".join(str(cell) for cell in row) 
                                               for row in rows])
                    else:
                        table_text = "\n".join([" | ".join(str(cell) for cell in row) 
                                               for row in rows])
                    
                    chunk["content"] = f"[표]\n{table_text}"
                else:
                    chunk["content"] = f"[표] {str(elem['content'])}"
                    
            elif elem["type"] == "image":
                # 이미지 설명 생성
                ocr_text = elem.get("metadata", {}).get("ocr_text", "")
                if ocr_text:
                    chunk["content"] = f"[이미지 내 텍스트] {ocr_text}"
                else:
                    chunk["content"] = f"[이미지] {elem['metadata'].get('format', 'unknown')} 형식, " \
                                     f"크기: {elem['metadata'].get('width', 'unknown')}x{elem['metadata'].get('height', 'unknown')}"
                
            elif elem["type"] == "equation":
                # 수식 처리
                content = elem["content"]
                if isinstance(content, dict):
                    latex = content.get("latex", "")
                    text = content.get("text", "")
                    chunk["content"] = f"[수식] {text or latex}"
                else:
                    chunk["content"] = f"[수식] {content}"
            
            # 관계 정보 추가
            if elem.get("relationships"):
                chunk["metadata"]["related_elements"] = elem["relationships"]
            
            if chunk["content"].strip():
                chunks.append(chunk)
        
        return chunks

# 편의 함수들
async def process_document_simple(file_path: Union[str, Path], 
                                 extract_images: bool = True, 
                                 extract_tables: bool = True) -> Dict[str, Any]:
    """간단한 문서 처리 함수"""
    processor = MultimodalDocumentProcessor()
    
    extract_options = {
        "extract_text": True,
        "extract_images": extract_images,
        "extract_tables": extract_tables
    }
    
    return await processor.process_document(file_path, extract_options, "structured_json")

async def process_documents_batch(file_paths: List[Union[str, Path]], 
                                max_workers: int = 4) -> Dict[str, Any]:
    """여러 문서 배치 처리"""
    import asyncio
    from concurrent.futures import ThreadPoolExecutor
    
    results = {
        "successful_documents": {},
        "failed_documents": {},
        "summary": {
            "total_files": len(file_paths),
            "successful": 0,
            "failed": 0
        }
    }
    
    processor = MultimodalDocumentProcessor()
    
    async def process_single(file_path):
        try:
            result = await processor.process_document(file_path)
            return str(file_path), result
        except Exception as e:
            return str(file_path), {"success": False, "error": str(e)}
    
    # 병렬 처리
    tasks = [process_single(fp) for fp in file_paths]
    processed_results = await asyncio.gather(*tasks)
    
    for file_path, result in processed_results:
        if result.get("success", False):
            results["successful_documents"][file_path] = result
            results["summary"]["successful"] += 1
        else:
            results["failed_documents"][file_path] = result
            results["summary"]["failed"] += 1
    
    return results